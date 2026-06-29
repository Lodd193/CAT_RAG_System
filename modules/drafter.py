import re
from datetime import datetime
from io import BytesIO

import anthropic
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

import config

_client = anthropic.Anthropic(api_key=config.ANTHROPIC_API_KEY)
_PROMPT_PATH = "prompts/drafting.txt"


def draft_document(doc_type: str, user_inputs: dict, live_context: str) -> str:
    with open(_PROMPT_PATH, encoding="utf-8") as f:
        prompt = f.read().format(
            doc_type=doc_type,
            user_inputs="\n".join(f"{k}: {v}" for k, v in user_inputs.items()),
            context=live_context,
        )
    response = _client.messages.create(
        model=config.CLAUDE_MODEL,
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )
    return response.content[0].text.strip()


# ── Docx helpers ───────────────────────────────────────────────────────────────

def _add_formatted_runs(paragraph, text: str) -> None:
    """Add runs to a paragraph, rendering **bold** and *italic* markers."""
    for match in re.finditer(r"\*\*(.+?)\*\*|\*(.+?)\*|([^*]+)", text):
        bold_text, italic_text, plain_text = match.group(1), match.group(2), match.group(3)
        if bold_text is not None:
            paragraph.add_run(bold_text).bold = True
        elif italic_text is not None:
            paragraph.add_run(italic_text).italic = True
        elif plain_text:
            paragraph.add_run(plain_text)


def _is_table_row(line: str) -> bool:
    return line.startswith("|") and line.endswith("|")


def _is_separator_row(line: str) -> bool:
    return _is_table_row(line) and all(c in "-|: " for c in line)


def _parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip("|").split("|")]


def _flush_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"

    # Stretch table to full text-area width
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "5000")   # 5000 twentieths of a percent = 100%
    tbl_w.set(qn("w:type"), "pct")

    for i, row_data in enumerate(rows):
        row = table.add_row()
        for j, cell_text in enumerate(row_data):
            if j >= ncols:
                break
            para = row.cells[j].paragraphs[0]
            _add_formatted_runs(para, cell_text)
            if i == 0:  # header row: force all text bold
                for run in para.runs:
                    run.bold = True

    # Space after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


# ── Public API ─────────────────────────────────────────────────────────────────

def create_docx(text: str, title: str = "") -> bytes:
    doc = Document()

    # A4 page, standard UK margins
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # Normal style: Calibri 11pt
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Document header block
    if title:
        heading = doc.add_heading(title, level=0)
        heading.paragraph_format.space_after = Pt(2)
        sub = doc.add_paragraph()
        sub.add_run("Clinical Administration Transformation (CAT) Programme").bold = True
        sub.add_run("  ·  University Hospitals Birmingham NHS Foundation Trust")
        sub.paragraph_format.space_after = Pt(2)
        date_p = doc.add_paragraph(datetime.now().strftime("%d %B %Y"))
        date_p.paragraph_format.space_after = Pt(14)

    table_buffer: list[list[str]] = []

    for line in text.splitlines():
        stripped = line.rstrip()

        if _is_table_row(stripped):
            if not _is_separator_row(stripped):
                table_buffer.append(_parse_table_row(stripped))
            continue  # consume separator rows silently

        # Flush any accumulated table before processing normal content
        if table_buffer:
            _flush_table(doc, table_buffer)
            table_buffer = []

        if stripped in ("---", "***", "___"):
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(4)
        elif stripped.startswith("### "):
            h = doc.add_heading(stripped[4:], level=3)
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(4)
        elif stripped.startswith("## "):
            h = doc.add_heading(stripped[3:], level=2)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
        elif stripped.startswith("# "):
            h = doc.add_heading(stripped[2:], level=1)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
        elif stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(p, stripped[2:])
            p.paragraph_format.space_after = Pt(3)
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_formatted_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            p.paragraph_format.space_after = Pt(3)
        elif stripped == "":
            pass  # blank lines: rely on paragraph space_after
        else:
            p = doc.add_paragraph()
            _add_formatted_runs(p, stripped)
            p.paragraph_format.space_after = Pt(6)

    # Flush any table at end of text
    if table_buffer:
        _flush_table(doc, table_buffer)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
